from fastapi import FastAPI, HTTPException, Response, WebSocket, BackgroundTasks, Body, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import requests
from typing import List
import logging
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import json
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__) 

app = FastAPI()

# CORS settings
origins = [
    "http://localhost",
    "http://localhost:8000",
    "http://127.0.0.1:8000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

class Query(BaseModel):
    search_term: str
    facets: List[str] = []

class ProcessedResult(BaseModel):
    source: str
    title: str
    snippet: str
    url: str

class ExportData(BaseModel):
    results: List[ProcessedResult]

@app.post("/search", response_model=List[ProcessedResult])
async def search(query: Query):
    search_term = query.search_term
    facets = query.facets
    results = []

    try:
        logger.info(f"Searching for term: {search_term} with facets: {facets}")

        google_results = process_google_search(search_term)
        results.extend(google_results)

        reddit_results = process_reddit_search(search_term)
        results.extend(reddit_results)

        processed_results = aggregate_and_process_results(results, facets)

        # Ensure snippet length constraints
        for result in processed_results:
            if len(result.snippet) > 250:
                result.snippet = result.snippet[:250] + "..."
            if len(result.snippet) < 210:
                result.snippet = result.snippet.ljust(210)

        return processed_results
    except Exception as e:
        logger.error(f"An error occurred: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/export_pdf")
async def export_pdf(data: ExportData):
    try:
        logger.info(f"Exporting PDF for {len(data.results)} results")
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, height - 50, "Search Results")
        c.setFont("Helvetica", 12)

        y_position = height - 100
        for result in data.results:
            if y_position < 100:
                c.showPage()
                y_position = height - 50

            c.setFont("Helvetica-Bold", 12)
            c.drawString(100, y_position, result.title)
            y_position -= 20
            c.setFont("Helvetica", 10)
            c.drawString(100, y_position, f"Source: {result.source}")
            y_position -= 15
            c.drawString(100, y_position, f"URL: {result.url}")
            y_position -= 15
            snippet = result.snippet[:250] + "..." if len(result.snippet) > 250 else result.snippet
            text_object = c.beginText(100, y_position)
            text_object.setFont("Helvetica", 9)
            for line in snippet.split('\n'):
                text_object.textLine(line)
            c.drawText(text_object)
            y_position -= 60

        c.save()
        buffer.seek(0)

        return Response(content=buffer.getvalue(), media_type="application/pdf", 
                        headers={"Content-Disposition": "attachment; filename=search_results.pdf"})
    except Exception as e:
        logger.error(f"Error exporting PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/export_docx")
async def export_docx(data: ExportData):
    try:
        logger.info(f"Exporting DOCX for {len(data.results)} results")
        document = Document()
        document.add_heading('Search Results', 0)

        for result in data.results:
            document.add_heading(result.title, level=1)
            document.add_paragraph(f"Source: {result.source}")
            document.add_paragraph(f"URL: {result.url}")
            document.add_paragraph(result.snippet)
            document.add_paragraph()  # Добавление пустого абзаца между результатами

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return Response(content=buffer.getvalue(), 
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={"Content-Disposition": "attachment; filename=search_results.docx"})
    except Exception as e:
        logger.error(f"Error exporting DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    while True:
        try:
            data = await websocket.receive_text()
            logger.info(f"Received WebSocket data: {data}")

            query = Query(**json.loads(data))
            results = await search(query)
            await websocket.send_json([result.dict() for result in results])

        except Exception as e:
            logger.error(f"WebSocket error: {str(e)}")
            await websocket.close()

def process_google_search(search_term: str):
    api_key = ""
    url = f"https://serpapi.com/search.json?q={search_term}&api_key={api_key}"
    response = requests.get(url)
    if response.status_code != 200:
        logger.error(f"Error fetching Google search results: {response.text}")
        raise HTTPException(status_code=500, detail="Error fetching Google search results")
    search_results = response.json().get('organic_results', [])

    results = []
    for result in search_results:
        results.append(ProcessedResult(
            source="Google",
            title=result.get("title"),
            snippet=result.get("snippet", ""),
            url=result.get("link")
        ))
    return results

def process_reddit_search(search_term: str):
    url = f"https://www.reddit.com/search.json?q={search_term}"
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        logger.error(f"Error fetching Reddit search results: {response.text}")
        raise HTTPException(status_code=500, detail="Error fetching Reddit search results")
    data = response.json()

    results = []
    for post in data.get('data', {}).get('children', []):
        post_data = post.get('data', {})
        title = post_data.get("title", "")
        snippet = post_data.get("selftext", "")
        permalink = post_data.get('permalink', '')
        url = f"https://www.reddit.com{permalink}"
        if title and permalink:
            results.append(ProcessedResult(
                source="Reddit",
                title=title,
                snippet=snippet,
                url=url
            ))
    return results

def aggregate_and_process_results(results: List[ProcessedResult], facets: List[str]):
    unique_urls = set()
    processed_results = []
    for result in results:
        if result.url not in unique_urls:
            unique_urls.add(result.url)
            if not facets or result.source in facets:
                processed_results.append(result)
    return processed_results

app.mount("/", StaticFiles(directory="static", html=True), name="static")